"""Report generation utilities for PDF and DOCX formats"""
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
from pathlib import Path
import asyncio
from datetime import datetime
import html


def escape_html(text: str) -> str:
    """Escape HTML special characters for safe use in reports"""
    if not text:
        return ""
    return html.escape(str(text))


async def generate_pdf_report(audit, results: List, reports_dir: Path) -> Path:
    """Generate a comprehensive PDF report"""
    
    def _generate():
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"audit_{audit.id}_{timestamp}.pdf"
        filepath = reports_dir / filename
        
        # Create PDF
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title with website name
        website_name = audit.website_url.split('//')[1].split('/')[0] if '//' in audit.website_url else audit.website_url
        story.append(Paragraph(f"SEO Audit Report for {website_name}", title_style))
        story.append(Paragraph("by MJ SEO", ParagraphStyle('subtitle', parent=styles['Normal'], fontSize=12, textColor=colors.HexColor('#6b7280'), alignment=TA_CENTER)))
        story.append(Spacer(1, 0.3*inch))
        
        # Personal introduction
        intro_text = f"""Welcome to your comprehensive SEO audit report! We've analyzed {audit.pages_crawled} pages from your website 
        and performed {audit.total_checks_run} detailed checks across 9 critical SEO categories. This report will show you exactly 
        what's working well, what needs attention, and most importantly - how to fix each issue with specific, actionable steps."""
        
        intro_style = ParagraphStyle('intro', parent=styles['BodyText'], fontSize=11, leading=16, spaceAfter=20)
        story.append(Paragraph(intro_text, intro_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Executive Summary
        story.append(Paragraph("📊 Your SEO Score Overview", heading_style))
        
        summary_data = [
            ["Website", audit.website_url],
            ["Audit Date", audit.created_at.strftime("%B %d, %Y at %H:%M")],
            ["Overall SEO Score", f"{audit.overall_score:.1f}/100"],
            ["Pages Analyzed", str(audit.pages_crawled)],
            ["Total Checks Performed", str(audit.total_checks_run)],
            ["✅ Checks Passed", str(audit.checks_passed)],
            ["❌ Issues Found", str(audit.checks_failed)],
            ["⚠️ Warnings", str(audit.checks_warning)],
        ]
        
        summary_table = Table(summary_data, colWidths=[2.2*inch, 3.8*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e0e7ff')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f8fafc')])
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Human-like Score Interpretation
        story.append(Paragraph("What Your Score Means", heading_style))
        score = audit.overall_score or 0
        if score >= 80:
            interpretation = f"""Great news! With a score of {score:.1f}/100, your website is performing well in terms of SEO. 
            You've got {audit.checks_passed} checks passing, which shows you're following SEO best practices. The items flagged 
            in this report are opportunities to make your already-good site even better. Focus on the high-impact issues first 
            for maximum results."""
        elif score >= 60:
            interpretation = f"""Your site scores {score:.1f}/100 - you're on the right track! You have {audit.checks_passed} 
            checks passing, which is a solid foundation. However, there are {audit.checks_failed} issues that could be holding 
            your rankings back. The good news? Most of these are fixable with the step-by-step solutions we've provided in this 
            report. Tackle the critical issues first, and you could see improvements within weeks."""
        elif score >= 40:
            interpretation = f"""With a score of {score:.1f}/100, your website needs some SEO attention. We found {audit.checks_failed} 
            issues across your {audit.pages_crawled} analyzed pages. Don't worry - this is actually common, and every issue in this 
            report comes with clear instructions on how to fix it. Start with the "Critical" and "High Impact" items first. These will 
            give you the biggest boost in search visibility."""
        else:
            interpretation = f"""Your current score of {score:.1f}/100 indicates significant SEO challenges that need immediate attention. 
            We've identified {audit.checks_failed} critical issues affecting your search performance. But here's the thing - this report 
            gives you a complete roadmap to fix everything. Follow the priority order we've laid out, starting with technical SEO fundamentals, 
            and you'll see steady improvement. Many sites have gone from similar scores to 70+ within 2-3 months by systematically addressing 
            these issues."""
        
        story.append(Paragraph(interpretation, intro_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Group results by category
        categories = {}
        for result in results:
            cat = result.category
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(result)
        
        # Detailed Results by Category
        story.append(PageBreak())
        story.append(Paragraph("Detailed Analysis", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        for category, cat_results in categories.items():
            story.append(Paragraph(category, heading_style))
            
            for result in cat_results[:10]:  # Limit to prevent huge PDFs
                # Check name
                check_heading = ParagraphStyle(
                    'CheckHeading',
                    parent=styles['Heading3'],
                    fontSize=12,
                    textColor=colors.HexColor('#4b5563'),
                    spaceAfter=6
                )
                story.append(Paragraph(f"• {result.check_name}", check_heading))
                
                # Status badge
                status_color = {
                    'pass': colors.green,
                    'fail': colors.red,
                    'warning': colors.orange,
                    'info': colors.blue
                }
                
                status_text = f"Status: {result.status.value.upper()}"
                status_para = Paragraph(
                    f"<font color='{status_color.get(result.status.value, colors.black)}'><b>{status_text}</b></font>",
                    styles['BodyText']
                )
                story.append(status_para)
                
                # Impact Score
                if result.impact_score:
                    story.append(Paragraph(f"<b>Impact Score:</b> {result.impact_score}/100", styles['BodyText']))
                
                # Current vs Recommended
                if result.current_value:
                    story.append(Paragraph(f"<b>Current:</b> {escape_html(result.current_value)}", styles['BodyText']))
                if result.recommended_value:
                    story.append(Paragraph(f"<b>Recommended:</b> {escape_html(result.recommended_value)}", styles['BodyText']))
                
                # Cons (issues)
                if result.cons:
                    story.append(Paragraph("<b>Issues:</b>", styles['BodyText']))
                    for con in result.cons[:3]:
                        story.append(Paragraph(f"  - {escape_html(con)}", styles['BodyText']))
                
                # Solution
                if result.solution:
                    story.append(Paragraph(f"<b>Solution:</b> {escape_html(result.solution[:300])}...", styles['BodyText']))
                
                story.append(Spacer(1, 0.15*inch))
            
            if len(cat_results) > 10:
                story.append(Paragraph(f"... and {len(cat_results) - 10} more checks in this category", styles['Italic']))
            
            story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(story)
        return filepath
    
    # Run in thread pool to avoid blocking
    return await asyncio.to_thread(_generate)


async def generate_docx_report(audit, results: List, reports_dir: Path) -> Path:
    """Generate a comprehensive DOCX report"""
    
    def _generate():
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"audit_{audit.id}_{timestamp}.docx"
        filepath = reports_dir / filename
        
        # Create DOCX
        doc = Document()
        
        # Title
        title = doc.add_heading('MJ SEO - Comprehensive SEO Audit Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        
        # Summary table
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        
        summary_items = [
            ('Website', audit.website_url),
            ('Audit Date', audit.created_at.strftime("%Y-%m-%d %H:%M")),
            ('Overall Score', f"{audit.overall_score:.1f}/100"),
            ('Pages Crawled', str(audit.pages_crawled)),
            ('Total Checks', str(audit.total_checks_run)),
            ('Passed', str(audit.checks_passed)),
            ('Failed', str(audit.checks_failed)),
            ('Warnings', str(audit.checks_warning))
        ]
        
        for idx, (label, value) in enumerate(summary_items):
            row = table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Make label bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Score Interpretation
        doc.add_heading('Score Interpretation', 2)
        score = audit.overall_score or 0
        if score >= 80:
            interpretation = "Excellent! Your site is well-optimized."
        elif score >= 60:
            interpretation = "Good, but there's room for improvement."
        elif score >= 40:
            interpretation = "Needs attention. Address critical issues first."
        else:
            interpretation = "Critical. Immediate action required."
        
        doc.add_paragraph(interpretation)
        
        # Group results by category
        categories = {}
        for result in results:
            cat = result.category
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(result)
        
        # Detailed Results
        doc.add_page_break()
        doc.add_heading('Detailed Analysis', 1)
        
        for category, cat_results in categories.items():
            doc.add_heading(category, 2)
            
            for result in cat_results:
                # Check name
                check_para = doc.add_paragraph()
                check_run = check_para.add_run(f"• {result.check_name}")
                check_run.bold = True
                check_run.font.size = Pt(12)
                
                # Status
                status_para = doc.add_paragraph()
                status_run = status_para.add_run(f"Status: {result.status.value.upper()}")
                status_run.bold = True
                
                # Color based on status
                status_colors = {
                    'pass': RGBColor(0, 128, 0),
                    'fail': RGBColor(255, 0, 0),
                    'warning': RGBColor(255, 165, 0),
                    'info': RGBColor(0, 0, 255)
                }
                status_run.font.color.rgb = status_colors.get(result.status.value, RGBColor(0, 0, 0))
                
                # Impact Score
                if result.impact_score:
                    impact_para = doc.add_paragraph()
                    impact_para.add_run('Impact Score: ').bold = True
                    impact_para.add_run(f"{result.impact_score}/100")
                
                # Current vs Recommended
                if result.current_value:
                    current_para = doc.add_paragraph()
                    current_para.add_run('Current: ').bold = True
                    current_para.add_run(result.current_value)
                
                if result.recommended_value:
                    rec_para = doc.add_paragraph()
                    rec_para.add_run('Recommended: ').bold = True
                    rec_para.add_run(result.recommended_value)
                
                # Ranking Impact
                if result.ranking_impact:
                    rank_para = doc.add_paragraph()
                    rank_para.add_run('Ranking Impact: ').bold = True
                    rank_para.add_run(result.ranking_impact)
                
                # Issues
                if result.cons:
                    doc.add_paragraph('Issues:').runs[0].bold = True
                    for con in result.cons:
                        doc.add_paragraph(f"  - {con}", style='List Bullet 2')
                
                # Solution
                if result.solution:
                    solution_para = doc.add_paragraph()
                    solution_para.add_run('Solution: ').bold = True
                    solution_para.add_run(result.solution)
                
                # Enhancements
                if result.enhancements:
                    doc.add_paragraph('Enhancement Suggestions:').runs[0].bold = True
                    for enhancement in result.enhancements[:5]:
                        doc.add_paragraph(f"  - {enhancement}", style='List Bullet 2')
                
                doc.add_paragraph()  # Spacing
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run('Generated by MJ SEO - Professional SEO Audit Platform')
        footer_run.italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        doc.save(str(filepath))
        return filepath
    
    # Run in thread pool to avoid blocking
    return await asyncio.to_thread(_generate)
