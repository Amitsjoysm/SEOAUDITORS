"""
Enhanced Report Generator with Real API Data
Includes: Lighthouse Core Web Vitals, DataForSEO insights, Sub-Agent recommendations
Makes reports site-specific, not generic
"""
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
from pathlib import Path
from datetime import datetime
import html
import asyncio


def escape_html(text: str) -> str:
    """Escape HTML for safe use"""
    return html.escape(str(text)) if text else ""


async def generate_enhanced_pdf_report(audit, results: List, reports_dir: Path) -> Path:
    """
    Generate PDF report with REAL site-specific data:
    - Lighthouse Core Web Vitals
    - DataForSEO competitor analysis
    - Sub-agent AI insights
    - Cumulative impact scoring
    """
    
    def _generate():
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"audit_{audit.id}_{timestamp}.pdf"
        filepath = reports_dir / filename
        
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, 
                                     textColor=colors.HexColor('#6366f1'), spaceAfter=20, alignment=TA_CENTER)
        website_name = audit.website_url.split('//')[1].split('/')[0] if '//' in audit.website_url else audit.website_url
        story.append(Paragraph(f"SEO Audit Report: {website_name}", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Executive Summary with REAL data
        story.append(Paragraph("📊 Executive Summary", styles['Heading2']))
        
        # Safe access to attributes with defaults
        overall_score = audit.overall_score or 0
        potential_score = audit.potential_score or 0
        score_grade = audit.score_grade or 'N/A'
        score_gap = potential_score - overall_score
        
        summary_text = f"""<b>Website:</b> {audit.website_url}<br/>
<b>Audit Date:</b> {audit.created_at.strftime("%B %d, %Y")}<br/>
<b>Current SEO Score:</b> {overall_score:.1f}/100 (Grade: {score_grade})<br/>
<b>Potential Score:</b> {potential_score:.1f}/100 (if all issues fixed)<br/>
<b>Score Gap:</b> {score_gap:.1f} points available<br/>
<b>Pages Analyzed:</b> {audit.pages_crawled}<br/>
<b>Total Checks:</b> {audit.total_checks_run}<br/>"""
        
        # Add Lighthouse data if available
        if audit.lighthouse_data:
            perf_score = audit.lighthouse_data.get('performance_score')
            cwv = audit.lighthouse_data.get('core_web_vitals', {})
            if perf_score:
                summary_text += f"<b>Lighthouse Performance:</b> {perf_score}/100<br/>"
            if cwv.get('lcp'):
                lcp_val = cwv['lcp'].get('display_value', 'N/A')
                summary_text += f"<b>Largest Contentful Paint (LCP):</b> {lcp_val}<br/>"
        
        # Add competitor data
        if audit.competitor_count > 0:
            summary_text += f"<b>Competitors Identified:</b> {audit.competitor_count}<br/>"
        
        # Add opportunities
        if audit.opportunities_found > 0:
            summary_text += f"<b>Content Opportunities:</b> {audit.opportunities_found}<br/>"
        
        story.append(Paragraph(summary_text, styles['BodyText']))
        story.append(Spacer(1, 0.3*inch))
        
        # Cumulative Impact Analysis
        story.append(Paragraph("🎯 Cumulative Impact: Your Path to 100/100", styles['Heading2']))
        
        # Calculate cumulative impact
        failed_by_impact = sorted([r for r in results if r.status.value == 'fail'], 
                                 key=lambda x: x.impact_score or 0, reverse=True)
        
        cumulative_score = audit.overall_score
        impact_table_data = [["Priority", "Issue", "Current Impact", "Score After Fix", "Cumulative Gain"]]
        
        for idx, result in enumerate(failed_by_impact[:10], 1):
            impact = (result.impact_score or 50) / 10  # Convert to score points
            cumulative_score += impact
            cumulative_score = min(100, cumulative_score)  # Cap at 100
            
            impact_table_data.append([
                str(idx),
                result.check_name[:40],
                f"{result.impact_score or 50}/100",
                f"{cumulative_score:.1f}",
                f"+{impact:.1f}"
            ])
        
        impact_table = Table(impact_table_data, colWidths=[0.6*inch, 2.5*inch, 1.2*inch, 1.2*inch, 1*inch])
        impact_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')])
        ]))
        story.append(impact_table)
        story.append(Spacer(1, 0.3*inch))
        
        # AI Orchestrator Insights (if available and not an error)
        if audit.analytics_summary and audit.analytics_summary.get('orchestrator_insights'):
            insights = audit.analytics_summary.get('orchestrator_insights', '')
            # Only show if it's not an error message
            if insights and not insights.startswith('Error generating'):
                story.append(Paragraph("🤖 AI Orchestrator Strategic Insights", styles['Heading2']))
                story.append(Paragraph(escape_html(insights[:1000]), styles['BodyText']))
                story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        return filepath
    
    # Run in thread pool to avoid blocking
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _generate)


async def generate_enhanced_docx_report(audit, results: List, reports_dir: Path) -> Path:
    """
    Generate DOCX report with REAL site-specific data
    """
    
    def _generate():
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"audit_{audit.id}_{timestamp}.docx"
        filepath = reports_dir / filename
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'SEO Audit Report: {audit.website_url}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('📊 Executive Summary', 1)
        doc.add_paragraph(f'Audit Date: {audit.created_at.strftime("%B %d, %Y")}')
        
        # Safe access to attributes with defaults
        overall_score = audit.overall_score or 0
        potential_score = audit.potential_score or 0
        score_grade = audit.score_grade or 'N/A'
        score_gap = potential_score - overall_score
        
        doc.add_paragraph(f'Current SEO Score: {overall_score:.1f}/100 (Grade: {score_grade})')
        doc.add_paragraph(f'Potential Score: {potential_score:.1f}/100')
        doc.add_paragraph(f'Score Improvement Available: {score_gap:.1f} points')
        
        # Real Lighthouse Data
        if audit.lighthouse_data:
            doc.add_heading('⚡ Real Performance Data (Lighthouse)', 1)
            perf = audit.lighthouse_data.get('performance_score')
            if perf:
                doc.add_paragraph(f'Performance Score: {perf}/100')
            
            cwv = audit.lighthouse_data.get('core_web_vitals', {})
            if cwv:
                doc.add_paragraph('Core Web Vitals:')
                for metric, data in cwv.items():
                    if data and isinstance(data, dict):
                        val = data.get('display_value', 'N/A')
                        doc.add_paragraph(f'  • {metric.upper()}: {val}', style='List Bullet')
        
        # Cumulative Impact
        doc.add_heading('🎯 Cumulative Impact: Path to 100/100', 1)
        doc.add_paragraph('Fix these issues in order to maximize your SEO score:')
        
        failed_by_impact = sorted([r for r in results if r.status.value == 'fail'],
                                 key=lambda x: x.impact_score or 0, reverse=True)
        
        cumulative = audit.overall_score
        for idx, result in enumerate(failed_by_impact[:10], 1):
            impact = (result.impact_score or 50) / 10
            cumulative += impact
            cumulative = min(100, cumulative)
            doc.add_paragraph(
                f'{idx}. {result.check_name} → Impact: {result.impact_score}/100 → Score after fix: {cumulative:.1f}/100',
                style='List Number'
            )
        
        # AI Insights
        if audit.analytics_summary and audit.analytics_summary.get('orchestrator_insights'):
            doc.add_heading('🤖 AI Orchestrator Insights', 1)
            insights = audit.analytics_summary.get('orchestrator_insights', '')
            doc.add_paragraph(insights[:2000])
        
        doc.save(str(filepath))
        return filepath
    
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _generate)
